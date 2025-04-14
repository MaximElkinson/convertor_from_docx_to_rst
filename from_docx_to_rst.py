from docx import Document as Document_for_reading
from datetime import datetime
import zipfile, pathlib

import queue
from spire.doc import *
from spire.doc.common import *

# Создайте объект документа
doc = Document()

# Загрузите файл Word
doc.LoadFromFile("admin.docx")

# Создайте объект очереди
nodes = queue.Queue()
nodes.put(doc)

# Создайте список
images = []

while nodes.qsize() > 0:
    node = nodes.get()

    # Переберите дочерние объекты в документе
    for i in range(node.ChildObjects.Count):
        child = node.ChildObjects.get_Item(i)

        # Определите, является ли дочерний объект изображением
        if child.DocumentObjectType == DocumentObjectType.Picture:
            picture = child if isinstance(child, DocPicture) else None
            dataBytes = picture.ImageBytes

            # Добавьте данные изображения в список
            images.append(dataBytes)

        elif isinstance(child, ICompositeObject):
            nodes.put(child if isinstance(child, ICompositeObject) else None)

# Переберите изображения в списке
name = f'{datetime.now()}'.split(' ')[0]
pathlib.Path(f'{name}').mkdir()
for i, item in enumerate(images):
    fileName = f"{name}/Image{i}.png"
    with open(fileName, 'wb') as imageFile:
        # Запишите изображение в указанное место
        imageFile.write(item)

def table_to_rst(table):
    rst_rows = []
    table = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    max_lens = [max(len(table[y][x]) for y in range(len(table))) for x in range(len(table[0]))]
    for row in table:
        rst_rows.append('+' + "+".join(['-' * i for i in max_lens]) + '+')
        rst_rows.append('|' + "|".join([row[i] + ' ' * (max_lens[i] - len(row[i])) for i in range(len(row))])  + '|')
    rst_rows.append('+' + "+".join(['-' * i for i in max_lens]) + '+')

    return '\n'.join(rst_rows)

def convert_word_table_to_rst(word_file, rst_file):
    # Открываем документ Word
    document = Document(word_file)
    rst_output = []

    # Перебираем все таблицы в документе
    for table in document.tables:
        rst_output.append(table_to_rst(table))
        rst_output.append("")  # Добавляем пустую строку между таблицами

    # Записываем результат в файл RST
    with open(rst_file, 'w') as f:
        f.write('\n'.join(rst_output))


def read_docx(file_path):
    # Открываем файл документа
    doc = Document_for_reading(file_path)
    text = []

    # Проходим по всем параграфам в документе
    for para in doc.paragraphs:
        # Проверяем стиль параграфа
        style = para.style.name

        # Выводим текст и соответствующий стиль
        text.append([style, para.text])
    return {'text': text, 'tables': doc.tables}

def extract_images(docx):
    # директория для извлечения
    ex_dir = pathlib.Path(f'pic_{docx}')
    if not ex_dir.is_dir():
        ex_dir.mkdir()

    with zipfile.ZipFile(docx) as zf:
        for name in zf.infolist():
            if name.filename.startswith('word/media'):
                # здесь можно задать другие параметры фильтрации,
                # например отобрать картинки с определенном именем,
                # расширением, размером `name.file_size` и т.д.
                if name.filename[-4:] != 'jpeg':
                    name.filename = name.filename.split('.')[0] + '.jpeg'
                zf.extract(name, ex_dir)





def write_to_rst(file_path, doc):
    text = doc['text']
    tables = doc['tables']

    with open(file_path, 'w', encoding='utf-8') as f:
        f.write('\n\n')
        pn = 1
        tn = 0
        rst = []
        for i in text:
            if i[0] == 'List Paragraph':
                rst.append(str(i[1]) + '\n')
                rst.append('=' * len(str(i[1])) + '\n\n')
            else:
                words = i[1].split('Рис.')
                rst.append(words[0] + '\n\n')
                for j in range(1, len(words)):
                    rst.append(f'.. image:: image{pn}.png' + '\n\n')
                    rst.append('Рис.' + words[j] + '\n\n')
                    pn += 1
        rst = ''.join(rst)
        rst = rst.split('Таб.')
        for i in range(len(rst)-1):
            rst[i] += '\n\n' + table_to_rst(tables[i]) + '\n\n'
        rst = 'Таб.'.join(rst)


        f.write(rst)

write_to_rst('index.rst', read_docx('admin.docx'), )

